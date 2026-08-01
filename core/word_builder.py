"""
word_builder.py - Image evidence Word document generation.
"""

from __future__ import annotations

import logging
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)


def build_image_docx(
    image_paths: list[Path],
    output_path: Path,
    title: str = "",
    show_number: bool = True,
) -> Path:
    """Build a Word document containing screenshots in order.

    The Word export intentionally does not run OCR or turn screenshots into text.
    It normalizes every image to PNG first because python-docx cannot insert
    every image format that the PDF path accepts, such as WebP.
    """
    if not image_paths:
        raise ValueError("image_paths cannot be empty")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    normalized_dir = output_path.parent / "_docx_images"
    normalized_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = section.page_height - section.top_margin - section.bottom_margin
    max_width_inches = usable_width / 914400
    max_height_inches = usable_height / 914400

    total = len(image_paths)
    for idx, img_path in enumerate(image_paths, start=1):
        img_path = Path(img_path)
        if show_number:
            number_p = doc.add_paragraph()
            _compact_paragraph(number_p)
            number_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            number_run = number_p.add_run(f"截图 {idx}")
            number_run.bold = True
            number_run.font.size = Pt(10)

        p = doc.add_paragraph()
        _compact_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        normalized = _normalize_for_docx(img_path, normalized_dir / f"image_{idx:04d}.png")
        height_reserve = 0.85 if show_number else 0.45
        width, height = _image_display_size(
            normalized,
            max_width_inches,
            max_height_inches - height_reserve,
        )
        run.add_picture(str(normalized), width=Inches(width), height=Inches(height))
        if idx < total:
            break_p = doc.add_page_break()
            _compact_paragraph(break_p)

    doc.save(str(output_path))
    logger.info("DOCX: %s (%.2f MB, %d images)",
                output_path.name,
                output_path.stat().st_size / 1024 / 1024,
                len(image_paths))
    return output_path


def _compact_paragraph(paragraph) -> None:
    """Remove Word's default paragraph spacing to avoid blank spill pages."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def _image_display_size(
    image_path: Path,
    max_width_inches: float,
    max_height_inches: float,
) -> tuple[float, float]:
    """Return a size that keeps the full screenshot inside one Word page."""
    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
        natural_width = max(width_px / 144, 0.1)
        natural_height = max(height_px / 144, 0.1)
        scale = min(
            max_width_inches / natural_width,
            max_height_inches / natural_height,
            1.0,
        )
        return max(0.1, natural_width * scale), max(0.1, natural_height * scale)
    except Exception:
        return max_width_inches, max_height_inches


def _normalize_for_docx(image_path: Path, output_path: Path) -> Path:
    """Convert any Pillow-readable image to a PNG file Word can embed."""
    try:
        with Image.open(image_path) as img:
            if img.mode in ("RGBA", "LA"):
                background = Image.new("RGB", img.size, "white")
                alpha = img.getchannel("A")
                background.paste(img.convert("RGB"), mask=alpha)
                normalized = background
            elif img.mode == "P":
                normalized = img.convert("RGBA")
                background = Image.new("RGB", normalized.size, "white")
                background.paste(normalized.convert("RGB"), mask=normalized.getchannel("A"))
                normalized = background
            else:
                normalized = img.convert("RGB")
            normalized.save(output_path, "PNG")
    except Exception as exc:
        raise RuntimeError(f"无法处理图片 {image_path.name}: {exc}") from exc
    return output_path
